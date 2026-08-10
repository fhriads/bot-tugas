import os
import shutil
import subprocess
import sys
from pathlib import Path


def convert_to_pdf(docx_path: str | Path, pdf_path: str | Path) -> str:
    """
    Converts a .docx document to a .pdf file.
    Primary engine: docx2pdf (Requires MS Word on Windows/macOS).
    Fallback engine: LibreOffice CLI ('soffice' / 'libreoffice').
    """
    docx_path = Path(docx_path).resolve()
    pdf_path = Path(pdf_path).resolve()
    pdf_path.parent.mkdir(parents=True, exist_ok=True)

    if not docx_path.exists():
        raise FileNotFoundError(f"Source DOCX file not found: {docx_path}")

    # Primary conversion attempt: docx2pdf
    try:
        from docx2pdf import convert
        print(f"[PDFConverter] Attempting conversion using docx2pdf...")
        convert(str(docx_path), str(pdf_path))
        if pdf_path.exists() and pdf_path.stat().st_size > 0:
            print(f"[PDFConverter] Successfully converted PDF via docx2pdf: {pdf_path}")
            return str(pdf_path)
    except Exception as e:
        print(f"[PDFConverter] docx2pdf conversion failed/unavailable ({e}). Trying fallback engine...")

    # Fallback conversion attempt: LibreOffice CLI
    libreoffice_bin = shutil.which("soffice") or shutil.which("libreoffice")
    if not libreoffice_bin and sys.platform == "win32":
        for possible_path in [
            Path(r"C:\Program Files\LibreOffice\program\soffice.exe"),
            Path(r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"),
        ]:
            if possible_path.exists():
                libreoffice_bin = str(possible_path)
                break

    if libreoffice_bin:
        try:
            print(f"[PDFConverter] Attempting conversion using LibreOffice ({libreoffice_bin})...")
            cmd = [
                libreoffice_bin,
                "--headless",
                "--convert-to", "pdf",
                "--outdir", str(pdf_path.parent),
                str(docx_path)
            ]
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
            if result.returncode == 0:
                # LibreOffice defaults output filename to same stem + .pdf
                generated_pdf = pdf_path.parent / f"{docx_path.stem}.pdf"
                if generated_pdf.exists():
                    if generated_pdf != pdf_path:
                        if pdf_path.exists():
                            pdf_path.unlink()
                        generated_pdf.rename(pdf_path)
                    print(f"[PDFConverter] Successfully converted PDF via LibreOffice: {pdf_path}")
                    return str(pdf_path)
            else:
                print(f"[PDFConverter] LibreOffice command error: {result.stderr}")
        except Exception as lo_err:
            print(f"[PDFConverter] LibreOffice conversion error: {lo_err}")

    # If both engines failed
    raise RuntimeError(
        "Gagal mengonversi file .docx ke .pdf!\n"
        "Pastikan Microsoft Word (dengan python package docx2pdf) atau LibreOffice terinstall di sistem Anda."
    )


def convert_scanned_pdf_to_docx(pdf_path: str | Path, output_docx_path: str | Path) -> str:
    """
    Converts a scanned PDF (or any PDF document) to a clean Word (.docx) document
    by rendering each page to a high-resolution image and building a structured DOCX.
    Guarantees zero missing images, accurate page sequence, no blank pages, and no overlapping elements.
    """
    import fitz  # PyMuPDF
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    pdf_path = Path(pdf_path).resolve()
    output_docx_path = Path(output_docx_path).resolve()
    output_docx_path.parent.mkdir(parents=True, exist_ok=True)

    if not pdf_path.exists():
        raise FileNotFoundError(f"File PDF tidak ditemukan: {pdf_path}")

    doc_pdf = fitz.open(str(pdf_path))
    doc_word = Document()

    # Set page margins to 0.5 inches for clean full-page rendering
    for section in doc_word.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    total_pages = len(doc_pdf)
    if total_pages == 0:
        doc_pdf.close()
        raise ValueError("File PDF kosong atau tidak memiliki halaman!")

    # Max boundaries for image to fit inside 0.5 in margins without triggering soft page break
    max_w = Inches(6.8)
    max_h = Inches(9.2)

    temp_images = []
    try:
        for page_num in range(total_pages):
            page = doc_pdf[page_num]
            # Render page at 200 DPI for high resolution without bloated file size
            pix = page.get_pixmap(dpi=200)

            temp_img_path = output_docx_path.parent / f"temp_conv_{output_docx_path.stem}_p{page_num}.png"
            pix.save(str(temp_img_path))
            temp_images.append(temp_img_path)

            if page_num > 0:
                doc_word.add_page_break()

            # Create zero-margin centered paragraph for image
            p = doc_word.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0

            # Calculate strict aspect ratio scaling to guarantee image fits within single page height
            img_w, img_h = pix.width, pix.height
            aspect_ratio = img_w / img_h
            calc_w = max_h * aspect_ratio

            if calc_w <= max_w:
                p.add_run().add_picture(str(temp_img_path), height=max_h)
            else:
                p.add_run().add_picture(str(temp_img_path), width=max_w)

        doc_word.save(str(output_docx_path))
        print(f"[PDFConverter] Successfully converted PDF to DOCX: {output_docx_path} ({total_pages} pages)")
        return str(output_docx_path)
    finally:
        doc_pdf.close()
        for img_p in temp_images:
            if img_p.exists():
                try:
                    img_p.unlink()
                except Exception:
                    pass


