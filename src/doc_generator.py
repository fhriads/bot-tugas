from pathlib import Path
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image

BASE_DIR = Path(__file__).resolve().parent.parent
DEFAULT_TEMPLATE_PATH = BASE_DIR / "templates" / "template_tugas.docx"
FALLBACK_TEMPLATE_PATH = BASE_DIR / "assets" / "template.docx"


def replace_text_in_paragraph(paragraph, replacements: dict):
    """
    Replaces dictionary key placeholders in a paragraph.
    Handles placeholders split across multiple run elements.
    """
    p_text = paragraph.text
    if not p_text:
        return

    needs_replacement = False
    for key in replacements:
        if key in p_text:
            needs_replacement = True
            break

    if not needs_replacement:
        return

    new_text = p_text
    for key, value in replacements.items():
        new_text = new_text.replace(key, str(value))

    if paragraph.runs:
        first_run = paragraph.runs[0]
        font_name = first_run.font.name
        font_size = first_run.font.size
        bold = first_run.bold
        italic = first_run.italic

        for run in paragraph.runs:
            run.text = ""

        first_run.text = new_text
        if font_name:
            first_run.font.name = font_name
        if font_size:
            first_run.font.size = font_size
        if bold is not None:
            first_run.bold = bold
        if italic is not None:
            first_run.italic = italic
    else:
        paragraph.text = new_text


def replace_placeholders_in_doc(doc, data_dict: dict):
    """
    Replaces placeholder variables across all paragraphs, tables, headers, and footers in the document.
    """
    mapping = {
        "{{NAMA}}": data_dict.get("nama", ""),
        "{{NIM}}": data_dict.get("nim", ""),
        "{{MATA_KULIAH}}": data_dict.get("mata_kuliah", data_dict.get("matkul", "")),
        "{{MATKUL}}": data_dict.get("mata_kuliah", data_dict.get("matkul", "")),
        "{{TANGGAL}}": data_dict.get("tanggal", ""),
    }

    for p in doc.paragraphs:
        replace_text_in_paragraph(p, mapping)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_text_in_paragraph(p, mapping)

    for section in doc.sections:
        for p in section.header.paragraphs:
            replace_text_in_paragraph(p, mapping)
        for p in section.footer.paragraphs:
            replace_text_in_paragraph(p, mapping)


def generate_docx(
    data_dict: dict,
    image_paths: list[str | Path],
    output_docx_path: str | Path,
    template_path: str | Path = None
) -> str:
    """
    Populates template.docx with profile metadata and appends processed scans.
    Strictly constrains Image 1 to max height 7.6 inches so it NEVER overflows Page 1.
    """
    output_docx_path = Path(output_docx_path)
    output_docx_path.parent.mkdir(parents=True, exist_ok=True)

    if template_path and Path(template_path).exists():
        tmpl = Path(template_path)
    elif DEFAULT_TEMPLATE_PATH.exists():
        tmpl = DEFAULT_TEMPLATE_PATH
    elif FALLBACK_TEMPLATE_PATH.exists():
        tmpl = FALLBACK_TEMPLATE_PATH
    else:
        from create_template import generate_default_template
        generate_default_template([DEFAULT_TEMPLATE_PATH])
        tmpl = DEFAULT_TEMPLATE_PATH

    doc = docx.Document(str(tmpl))

    # 1. Fill Placeholders
    replace_placeholders_in_doc(doc, data_dict)

    # 2. Append Processed Image Pages with Strict Zero-Overflow Calculation
    for idx, img_p in enumerate(image_paths):
        img_path_obj = Path(img_p)
        if not img_path_obj.exists():
            print(f"[DocGenerator] Warning: Image file not found: {img_p}")
            continue

        if idx > 0:
            doc.add_page_break()

        # Create compact paragraph for image
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0

        # Read actual image dimensions
        with Image.open(img_path_obj) as im:
            img_w, img_h = im.size

        # Define maximum boundaries
        if idx == 0:
            # Page 1: Strict height limit of 7.6 inches to guarantee Image 1 stays on Page 1 with header
            max_w = Inches(6.8)
            max_h = Inches(7.6)
        else:
            # Page 2+: Full page printable height limit of 9.4 inches
            max_w = Inches(6.8)
            max_h = Inches(9.4)

        # Calculate exact aspect ratio scaling
        aspect_ratio = img_w / img_h
        calc_w = max_h * aspect_ratio

        if calc_w <= max_w:
            p.add_run().add_picture(str(img_path_obj), height=max_h)
        else:
            p.add_run().add_picture(str(img_path_obj), width=max_w)

    # 3. Save Output Document
    doc.save(str(output_docx_path))
    print(f"[DocGenerator] Successfully generated: {output_docx_path}")
    return str(output_docx_path)
