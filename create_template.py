import os
from pathlib import Path
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

BASE_DIR = Path(__file__).resolve().parent
TEMPLATES_DIR = BASE_DIR / "templates"
ASSETS_DIR = BASE_DIR / "assets"


def set_cell_background(cell, fill_hex: str):
    """Sets background color of a docx table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)


def set_cell_margins(cell, top=20, bottom=20, left=60, right=60):
    """Sets tight internal cell padding (dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)


def generate_default_template(output_paths: list[Path]):
    """Generates a clean, bulletproof compact header docx template."""
    doc = docx.Document()

    # Set page margins strictly to 0.4 inch (top/bottom) and 0.5 inch (left/right)
    for section in doc.sections:
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.4)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # Document Header Title - Clean Black Color
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(2)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    title_run = title_p.add_run("LEMBAR PENGUMPULAN TUGAS")
    title_run.font.name = "Arial"
    title_run.font.size = Pt(12)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 0, 0)

    # Sleek 2-row x 4-column compact horizontal grid table
    table = doc.add_table(rows=2, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    grid = [
        [("Nama", "{{NAMA}}"), ("NIM", "{{NIM}}")],
        [("Matkul", "{{MATA_KULIAH}}"), ("Tanggal", "{{TANGGAL}}")]
    ]

    col_widths = [Inches(1.0), Inches(2.5), Inches(1.0), Inches(2.5)]

    for row_idx, row_items in enumerate(grid):
        row = table.rows[row_idx]
        col_offset = 0
        for label, placeholder in row_items:
            # Label cell
            c_lbl = row.cells[col_offset]
            c_lbl.width = col_widths[col_offset]
            set_cell_background(c_lbl, "F0F4F8")
            set_cell_margins(c_lbl, top=20, bottom=20, left=60, right=60)
            p_lbl = c_lbl.paragraphs[0]
            p_lbl.paragraph_format.space_before = Pt(0)
            p_lbl.paragraph_format.space_after = Pt(0)
            r_lbl = p_lbl.add_run(label)
            r_lbl.font.name = "Arial"
            r_lbl.font.size = Pt(8.5)
            r_lbl.font.bold = True

            # Value cell
            c_val = row.cells[col_offset + 1]
            c_val.width = col_widths[col_offset + 1]
            set_cell_margins(c_val, top=20, bottom=20, left=60, right=60)
            p_val = c_val.paragraphs[0]
            p_val.paragraph_format.space_before = Pt(0)
            p_val.paragraph_format.space_after = Pt(0)
            r_val = p_val.add_run(placeholder)
            r_val.font.name = "Arial"
            r_val.font.size = Pt(8.5)

            col_offset += 2

    # Spacing after table before images
    sep_p = doc.add_paragraph()
    sep_p.paragraph_format.space_before = Pt(2)
    sep_p.paragraph_format.space_after = Pt(2)

    # Save to target output paths
    for p in output_paths:
        p.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(p))
        print(f"[TemplateGenerator] Updated template at: {p}")


if __name__ == "__main__":
    targets = [
        TEMPLATES_DIR / "template_tugas.docx",
        ASSETS_DIR / "template.docx"
    ]
    generate_default_template(targets)
